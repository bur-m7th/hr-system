import sys
import json
from docx import Document

def main():
    if len(sys.argv) < 4:
        print("Usage: python generate_doc.py <template> <output> <json_file>")
        sys.exit(1)

    template_path = sys.argv[1]
    output_path = sys.argv[2]
    json_file = sys.argv[3]

    # READ JSON FROM FILE
    try:
        with open(json_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        print(f"✓ Successfully loaded JSON from {json_file}")
        print(f"Data: {json.dumps(data, indent=2)}")
    except Exception as e:
        print(f"ERROR reading JSON file: {e}")
        sys.exit(1)

    # Clean Data Helpers
    def safe_str(val):
        if val is None: 
            return ""
        s = str(val).strip()
        if s == ".": 
            return "" 
        return s

    def fmt_num(val):
        try:
            return "{:,.3f}".format(float(val))
        except:
            return "0.000"

    # Define Replacements
    replacements = {
        "{{EMPLOYEE_NAME}}": safe_str(data.get("employee_name")),
        "{{EMPLOYEE_ID}}": safe_str(data.get("employee_id")),
        "{{POSITION}}": safe_str(data.get("position")),
        "{{DEPARTMENT}}": safe_str(data.get("department")),
        "{{EMAIL}}": safe_str(data.get("email")),
        "{{PHONE}}": safe_str(data.get("phone")),
        "{{ADDRESS}}": safe_str(data.get("address")),
        "{{NATIONAL_ID}}": safe_str(data.get("national_id")),
        "{{BASE_SALARY}}": fmt_num(data.get("base_salary", 0)),
        "{{BONUS}}": fmt_num(data.get("bonus", 0)),
        "{{DEDUCTIONS}}": fmt_num(data.get("deductions", 0)),
        "{{NET_SALARY}}": fmt_num(data.get("net_salary", 0)),
        "{{PAY_PERIOD}}": safe_str(data.get("pay_period")),
        "{{CONTRACT_START}}": safe_str(data.get("contract_start")),
        "{{CONTRACT_END}}": safe_str(data.get("contract_end")),
        "{{GENERATED_DATE}}": safe_str(data.get("generated_date")),
    }

    print("\n--- STARTING REPLACEMENT ---")
    print(f"Template: {template_path}")
    print(f"Output: {output_path}")
    
    try:
        doc = Document(template_path)
        found_any = False
        replacement_count = 0

        def replace_in_paragraph(paragraph):
            """Replace placeholders in paragraph"""
            nonlocal found_any, replacement_count
            
            for key, value in replacements.items():
                if key in paragraph.text:
                    found_any = True
                    # Replace in each run
                    for run in paragraph.runs:
                        if key in run.text:
                            print(f"  ✓ Replacing {key} -> '{value}'")
                            run.text = run.text.replace(key, value)
                            replacement_count += 1

        # Scan all document parts
        print("\nScanning document...")
        
        for p in doc.paragraphs:
            replace_in_paragraph(p)
            
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace_in_paragraph(p)
        
        for section in doc.sections:
            # Headers
            for header in [section.header, section.first_page_header, section.even_page_header]:
                for p in header.paragraphs:
                    replace_in_paragraph(p)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                replace_in_paragraph(p)
            
            # Footers
            for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
                for p in footer.paragraphs:
                    replace_in_paragraph(p)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                replace_in_paragraph(p)

        if not found_any:
            print("\n⚠ WARNING: No placeholders found in the document!")
            print("Template may be corrupted or placeholders are missing.")
            print("\nExpected placeholders:")
            for key in replacements.keys():
                print(f"  - {key}")
        else:
            print(f"\n✓ Made {replacement_count} replacements successfully")

        doc.save(output_path)
        print(f"✓ Document saved to: {output_path}")

    except Exception as e:
        print(f"\n❌ ERROR: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

if __name__ == "__main__":
    main()